#!/usr/bin/env python3
"""
Module xử lý file PDF và Word cho CV Job Matching API
Hỗ trợ extract text từ PDF, DOCX, DOC files
"""

import os
import re
from typing import Optional, Union
from pathlib import Path
import logging

# PDF processing
try:
    import PyPDF2
    import fitz  # PyMuPDF - alternative PDF reader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word processing  
try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Alternative Word processing
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

# For .doc files (older format)
try:
    import win32com.client
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Class để xử lý các loại document khác nhau
    """
    
    def __init__(self):
        self.supported_formats = []
        if PDF_AVAILABLE:
            self.supported_formats.extend(['.pdf'])
        if DOCX_AVAILABLE or DOCX2TXT_AVAILABLE:
            self.supported_formats.extend(['.docx'])
        if WIN32_AVAILABLE:
            self.supported_formats.extend(['.doc'])
    
    def get_supported_formats(self) -> list:
        """Trả về list các format được hỗ trợ"""
        return self.supported_formats
    
    def extract_text_from_pdf(self, file_path: str, method: str = "pypdf2") -> str:
        """
        Extract text từ PDF file
        
        Args:
            file_path: Đường dẫn đến file PDF
            method: "pypdf2" hoặc "pymupdf"
        
        Returns:
            str: Text content của PDF
        """
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing libraries not available. Install: pip install PyPDF2 PyMuPDF")
        
        text = ""
        
        try:
            if method == "pypdf2":
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text += page.extract_text() + "\n"
            
            elif method == "pymupdf":
                doc = fitz.open(file_path)
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text += page.get_text() + "\n"
                doc.close()
            
            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            if method == "pypdf2" and 'fitz' in globals():
                logger.info("Trying PyMuPDF as fallback...")
                return self.extract_text_from_pdf(file_path, "pymupdf")
            raise e
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text từ DOCX file
        
        Args:
            file_path: Đường dẫn đến file DOCX
            
        Returns:
            str: Text content của DOCX
        """
        text = ""
        
        try:
            # Method 1: python-docx
            if DOCX_AVAILABLE:
                doc = Document(file_path)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                    text += "\n"
            
            # Method 2: docx2txt (fallback)
            elif DOCX2TXT_AVAILABLE:
                text = python_docx2txt.process(file_path)
            
            else:
                raise ImportError("DOCX processing libraries not available. Install: pip install python-docx")
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            # Try fallback method
            if DOCX_AVAILABLE and DOCX2TXT_AVAILABLE:
                try:
                    text = python_docx2txt.process(file_path)
                    return text.strip()
                except:
                    pass
            raise e
    
    def extract_text_from_doc(self, file_path: str) -> str:
        """
        Extract text từ DOC file (older Word format)
        Chỉ hoạt động trên Windows với Microsoft Word installed
        
        Args:
            file_path: Đường dẫn đến file DOC
            
        Returns:
            str: Text content của DOC
        """
        if not WIN32_AVAILABLE:
            raise ImportError("Win32 COM client not available. This only works on Windows with MS Word installed.")
        
        try:
            # Convert to absolute path
            file_path = os.path.abspath(file_path)
            
            # Create Word application
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            
            # Open document
            doc = word_app.Documents.Open(file_path)
            
            # Extract text
            text = doc.Content.Text
            
            # Close document and application
            doc.Close()
            word_app.Quit()
            
            logger.info(f"Successfully extracted {len(text)} characters from DOC")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOC {file_path}: {e}")
            raise e
    
    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text từ file dựa trên extension
        
        Args:
            file_path: Đường dẫn đến file
            
        Returns:
            str: Text content của file
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_extension == '.doc':
            return self.extract_text_from_doc(file_path)
        elif file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def clean_extracted_text(self, text: str) -> str:
        """
        Clean và normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[^\w\s\-\.\,\;\:\!\?\(\)]', ' ', text)
        
        # Remove extra spaces
        text = ' '.join(text.split())
        
        # Remove very short lines (likely formatting artifacts)
        lines = text.split('\n')
        cleaned_lines = [line.strip() for line in lines if len(line.strip()) > 3]
        
        return '\n'.join(cleaned_lines)
    
    def process_document(self, file_path: str, clean: bool = True) -> dict:
        """
        Process document và return thông tin chi tiết
        
        Args:
            file_path: Đường dẫn đến file
            clean: Có clean text hay không
            
        Returns:
            dict: Thông tin về document và text content
        """
        try:
            # Extract text
            raw_text = self.extract_text_from_file(file_path)
            
            # Clean text if requested
            cleaned_text = self.clean_extracted_text(raw_text) if clean else raw_text
            
            # File info
            file_stat = os.stat(file_path)
            file_info = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'file_size': file_stat.st_size,
                'file_extension': Path(file_path).suffix.lower(),
                'raw_text_length': len(raw_text),
                'cleaned_text_length': len(cleaned_text),
                'word_count': len(cleaned_text.split()),
                'processing_success': True,
                'raw_text': raw_text,
                'cleaned_text': cleaned_text
            }
            
            return file_info
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'processing_success': False,
                'error': str(e),
                'raw_text': '',
                'cleaned_text': ''
            }

# Convenience functions
def extract_text(file_path: str, clean: bool = True) -> str:
    """
    Convenience function để extract text từ file
    
    Args:
        file_path: Đường dẫn đến file
        clean: Có clean text hay không
        
    Returns:
        str: Text content
    """
    processor = DocumentProcessor()
    result = processor.process_document(file_path, clean)
    
    if not result['processing_success']:
        raise Exception(result['error'])
    
    return result['cleaned_text'] if clean else result['raw_text']

def get_supported_formats() -> list:
    """
    Get list các file format được hỗ trợ
    
    Returns:
        list: List các extension được hỗ trợ
    """
    processor = DocumentProcessor()
    return processor.get_supported_formats()

def check_dependencies() -> dict:
    """
    Check xem các dependency có available không
    
    Returns:
        dict: Status của các library
    """
    return {
        'pdf_support': PDF_AVAILABLE,
        'docx_support': DOCX_AVAILABLE or DOCX2TXT_AVAILABLE,
        'doc_support': WIN32_AVAILABLE,
        'supported_formats': get_supported_formats()
    }

# Example usage
if __name__ == "__main__":
    processor = DocumentProcessor()
    
    print("📄 Document Processor Test")
    print("=" * 30)
    
    # Check dependencies
    deps = check_dependencies()
    print("Dependencies:")
    for key, value in deps.items():
        print(f"  {key}: {value}")
    
    print(f"\nSupported formats: {get_supported_formats()}")
    
    # Test with sample files if available
    test_files = [
        "sample.pdf",
        "sample.docx", 
        "sample.doc",
        "sample.txt"
    ]
    
    for test_file in test_files:
        if os.path.exists(test_file):
            try:
                result = processor.process_document(test_file)
                print(f"\n✅ {test_file}:")
                print(f"  Word count: {result['word_count']}")
                print(f"  Text preview: {result['cleaned_text'][:100]}...")
            except Exception as e:
                print(f"\n❌ {test_file}: {e}")
